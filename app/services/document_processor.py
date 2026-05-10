#!/usr/bin/env python3
"""
文档处理器
用于处理多种格式的文档文件（.docx, .pdf, .doc, .txt等）
"""

import os
import re
import json
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging
from datetime import datetime

# 尝试导入文档处理库
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx 库未安装，无法处理 .docx 文件")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 库未安装，无法处理 .pdf 文件")

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("pandas 库未安装，无法处理 .xlsx 文件")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, data_cleaner=None):
        """初始化文档处理器"""
        self.data_cleaner = data_cleaner
        self.supported_formats = {
            '.docx': self._process_docx,
            '.pdf': self._process_pdf,
            '.doc': self._process_doc,
            '.txt': self._process_txt,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
        }
        
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        处理单个文档文件
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            处理后的文档数据
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                logger.error(f"文件不存在: {file_path}")
                return {"error": "文件不存在", "file_path": str(file_path)}
            
            # 获取文件扩展名
            ext = file_path.suffix.lower()
            
            # 检查是否支持该格式
            if ext not in self.supported_formats:
                logger.warning(f"不支持的文件格式: {ext}")
                return {"error": f"不支持的文件格式: {ext}", "file_path": str(file_path)}
            
            # 检查必要的库是否可用
            if ext == '.docx' and not DOCX_AVAILABLE:
                return {"error": "python-docx 库未安装", "file_path": str(file_path)}
            elif ext == '.pdf' and not PDF_AVAILABLE:
                return {"error": "PyPDF2 库未安装", "file_path": str(file_path)}
            elif ext == '.xlsx' and not EXCEL_AVAILABLE:
                return {"error": "pandas 库未安装", "file_path": str(file_path)}
            
            # 调用对应的处理函数
            processor_func = self.supported_formats[ext]
            result = processor_func(file_path)
            
            # 添加元数据
            result["metadata"] = {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "last_modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                "file_extension": ext,
                "processing_date": datetime.now().isoformat()
            }
            
            # 如果提供了数据清洗器，则清洗文本
            if self.data_cleaner and "content" in result:
                result["cleaned_content"] = self.data_cleaner.clean_document_text(result["content"])
            
            logger.info(f"成功处理文档: {file_path.name}")
            return result
            
        except Exception as e:
            logger.error(f"处理文档失败 {file_path}: {str(e)}")
            return {"error": str(e), "file_path": str(file_path)}
    
    def process_directory(self, directory_path: str, recursive: bool = True) -> List[Dict[str, Any]]:
        """
        处理目录中的所有文档
        
        Args:
            directory_path: 目录路径
            recursive: 是否递归处理子目录
            
        Returns:
            所有处理后的文档数据列表
        """
        results = []
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            logger.error(f"目录不存在: {directory_path}")
            return results
        
        # 遍历文件
        if recursive:
            file_iterator = directory_path.rglob("*")
        else:
            file_iterator = directory_path.glob("*")
        
        for file_path in file_iterator:
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.supported_formats:
                    result = self.process_document(str(file_path))
                    if "error" not in result:
                        results.append(result)
        
        logger.info(f"处理完成，共处理 {len(results)} 个文档")
        return results
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """处理 .docx 文件"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        content.append(" | ".join(row_text))
            
            full_text = "\n".join(content)
            
            return {
                "content": full_text,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "format": "docx"
            }
        except Exception as e:
            logger.error(f"处理 .docx 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "docx"}
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """处理 .pdf 文件"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = []
                
                # 提取每一页的文本
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
                
                full_text = "\n".join(content)
                
                return {
                    "content": full_text,
                    "page_count": len(pdf_reader.pages),
                    "format": "pdf"
                }
        except Exception as e:
            logger.error(f"处理 .pdf 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "pdf"}
    
    def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """处理 .doc 文件（旧版Word文档）"""
        # .doc 文件处理比较复杂，这里先返回简单文本
        try:
            # 尝试使用文本模式读取
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            return {
                "content": content,
                "format": "doc",
                "note": "使用文本模式读取，可能丢失格式信息"
            }
        except Exception as e:
            logger.error(f"处理 .doc 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "doc"}
    
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """处理 .txt 文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "content": content,
                "format": "txt"
            }
        except Exception as e:
            logger.error(f"处理 .txt 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "txt"}
    
    def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """处理 .json 文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # 如果是列表，转换为文本
            if isinstance(data, list):
                content = "\n".join([str(item) for item in data])
            elif isinstance(data, dict):
                content = json.dumps(data, ensure_ascii=False, indent=2)
            else:
                content = str(data)
            
            return {
                "content": content,
                "data": data,
                "format": "json"
            }
        except Exception as e:
            logger.error(f"处理 .json 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "json"}
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """处理 .csv 文件"""
        try:
            import csv
            content = []
            
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    content.append(", ".join(row))
            
            full_text = "\n".join(content)
            
            return {
                "content": full_text,
                "format": "csv"
            }
        except Exception as e:
            logger.error(f"处理 .csv 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "csv"}
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """处理 .xlsx 文件"""
        try:
            df = pd.read_excel(file_path)
            content = []
            
            # 添加列名
            content.append(", ".join(df.columns.tolist()))
            
            # 添加数据行
            for _, row in df.iterrows():
                content.append(", ".join([str(val) for val in row]))
            
            full_text = "\n".join(content)
            
            return {
                "content": full_text,
                "row_count": len(df),
                "column_count": len(df.columns),
                "format": "excel"
            }
        except Exception as e:
            logger.error(f"处理 .xlsx 文件失败 {file_path}: {str(e)}")
            return {"error": str(e), "format": "excel"}
    
    def categorize_document(self, document_data: Dict[str, Any]) -> str:
        """
        对文档进行分类
        
        Args:
            document_data: 文档数据
            
        Returns:
            文档分类
        """
        content = document_data.get("content", "").lower()
        file_name = document_data.get("metadata", {}).get("file_name", "").lower()
        
        # 根据关键词分类
        categories = {
            "regulation": ["条例", "规定", "法规", "法律", "法", "规章", "标准"],
            "case": ["案例", "判决", "裁决", "仲裁", "纠纷", "争议", "诉讼"],
            "template": ["模板", "范本", "示例", "样本", "格式"],
            "guideline": ["指南", "指导", "意见", "建议", "说明", "解读"],
            "analysis": ["分析", "解析", "研究", "探讨", "评述"],
            "other": []
        }
        
        # 检查文件名和内容中的关键词
        for category, keywords in categories.items():
            for keyword in keywords:
                if keyword in content or keyword in file_name:
                    return category
        
        return "other"
    
    def extract_key_info(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        提取文档关键信息
        
        Args:
            document_data: 文档数据
            
        Returns:
            关键信息字典
        """
        content = document_data.get("content", "")
        
        # 提取日期
        date_pattern = r'\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?'
        dates = re.findall(date_pattern, content)
        
        # 提取法律条款编号
        law_pattern = r'第[零一二三四五六七八九十百千万\d]+条'
        law_articles = re.findall(law_pattern, content)
        
        # 提取金额
        amount_pattern = r'[¥￥]?\s*\d+(?:,\d{3})*(?:\.\d+)?\s*(?:元|万元|亿元)?'
        amounts = re.findall(amount_pattern, content)
        
        # 提取百分比
        percent_pattern = r'\d+(?:\.\d+)?%'
        percents = re.findall(percent_pattern, content)
        
        return {
            "dates": list(set(dates))[:10],  # 去重并限制数量
            "law_articles": list(set(law_articles))[:20],
            "amounts": list(set(amounts))[:10],
            "percents": list(set(percents))[:10],
            "word_count": len(content),
            "character_count": len(content.replace(" ", "")),
            "line_count": content.count("\n") + 1
        }

def main() -> None:
    """测试文档处理器"""
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python document_processor.py <文件或目录路径>")
        return
    
    path = sys.argv[1]
    processor = DocumentProcessor()
    
    if os.path.isfile(path):
        result = processor.process_document(path)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif os.path.isdir(path):
        results = processor.process_directory(path)
        print(f"共处理 {len(results)} 个文档")
        
        # 统计分类
        categories = {}
        for result in results:
            category = processor.categorize_document(result)
            categories[category] = categories.get(category, 0) + 1
        
        print("文档分类统计:")
        for category, count in categories.items():
            print(f"  {category}: {count}")
    else:
        print(f"路径不存在: {path}")

if __name__ == "__main__":
    main()
