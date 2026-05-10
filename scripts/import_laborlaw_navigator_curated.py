#!/usr/bin/env python3
"""
Curate useful LaborLawNavigator files and import them into Chroma.

The assistant is scoped to Chongqing labor arbitration. This importer keeps
national labor-law rules, arbitration workflow materials, useful templates,
and generally applicable labor-dispute cases. It excludes obvious non-Chongqing
local regulations/cases and repository metadata.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import chromadb
from chromadb.config import Settings

from app.services.local_embedding import embed_texts

try:
    import docx
except Exception:
    docx = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import pandas as pd
except Exception:
    pd = None


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATASET_DIR = PROJECT_ROOT / "LaborLawNavigator"
OUTPUT_DIR = PROJECT_ROOT / "data" / "external_datasets"
MANIFEST_PATH = OUTPUT_DIR / "laborlaw_navigator_curated_manifest.json"
CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", str(PROJECT_ROOT / "data" / "chroma_db"))
COLLECTION_NAME = os.getenv("LABORLAW_NAVIGATOR_COLLECTION", "laborlaw_navigator_curated")

SUPPORTED_SUFFIXES = {".docx", ".doc", ".pdf", ".xlsx", ".csv", ".txt", ".json"}
EXCLUDE_NAME_KEYWORDS = {
    ".git",
    "上海",
    "北京",
    "京市",
    "顺义",
    "海南",
    "琼劳",
    "国际公约",
    "1958年消除就业和职业歧视公约",
    "准予就业最低年龄公约",
    "dataset_infos",
    "README",
    ".gitattributes",
}
USEFUL_KEYWORDS = {
    "劳动法",
    "劳动合同法",
    "劳动争议",
    "劳动纠纷",
    "仲裁",
    "工资",
    "报酬",
    "加班",
    "工伤",
    "社保",
    "解除",
    "终止",
    "竞业",
    "年休假",
    "女职工",
    "工会",
    "职业病",
    "劳动保障监察",
    "最低工资",
    "案例",
    "判决",
    "裁决",
    "文书",
    "申请书",
    "通知书",
    "证据",
    "相关法规",
    "劳动者",
}
CATEGORY_RULES = [
    ("case", ["案例", "判决", "裁决", "裁定"]),
    ("template", ["模板", "范文", "申请书", "起诉书", "答辩状", "通知书", "文书", "协议书"]),
    ("guide", ["咨询", "问答", "手册", "解析", "解读", "指南"]),
    ("regulation", ["法", "条例", "规定", "办法", "实施细则", "解释", "暂行规定"]),
]


def main() -> None:
    if not DATASET_DIR.exists():
        raise FileNotFoundError(f"Dataset directory not found: {DATASET_DIR}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    records, skipped = build_records()
    documents = [record["document"] for record in records]
    embeddings = embed_texts(documents)

    client = chromadb.PersistentClient(
        path=os.path.abspath(CHROMA_PERSIST_DIR),
        settings=Settings(anonymized_telemetry=False),
    )
    try:
        client.delete_collection(name=COLLECTION_NAME)
    except Exception:
        pass
    collection = client.get_or_create_collection(name=COLLECTION_NAME)
    if records:
        collection.upsert(
            ids=[record["id"] for record in records],
            documents=documents,
            metadatas=[record["metadata"] for record in records],
            embeddings=embeddings,
        )

    manifest = {
        "collection": COLLECTION_NAME,
        "persist_dir": os.path.abspath(CHROMA_PERSIST_DIR),
        "dataset_dir": str(DATASET_DIR),
        "selected_files": sorted({record["metadata"]["source_file"] for record in records}),
        "selected_file_count": len({record["metadata"]["source_file"] for record in records}),
        "chunk_count": len(records),
        "skipped": skipped,
        "selection_policy": {
            "kept": "national labor-law rules, arbitration workflow materials, templates, and generally applicable labor-dispute cases",
            "excluded": "non-Chongqing local regulations/cases, repository metadata, unsupported files, and files with too little extractable text",
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps({
        "collection": COLLECTION_NAME,
        "imported_chunks": len(records),
        "selected_files": manifest["selected_file_count"],
        "total_documents": collection.count(),
        "manifest": str(MANIFEST_PATH),
    }, ensure_ascii=False, indent=2))


def build_records() -> Tuple[List[Dict[str, Any]], List[Dict[str, str]]]:
    records: List[Dict[str, Any]] = []
    skipped: List[Dict[str, str]] = []
    for file_path in sorted(DATASET_DIR.rglob("*")):
        if not file_path.is_file() or ".git" in file_path.parts:
            continue
        decision, reason = should_include_file(file_path)
        if not decision:
            skipped.append({"file": file_path.name, "reason": reason})
            continue
        text = extract_text(file_path)
        text = clean_text(text)
        if len(text) < 120:
            skipped.append({"file": file_path.name, "reason": "extractable text too short"})
            continue
        category = categorize(file_path.name, text[:1000])
        for chunk_index, chunk in enumerate(chunk_text(text)):
            document = format_document(file_path.name, category, chunk_index, chunk)
            records.append({
                "id": stable_id(file_path.name, chunk_index),
                "document": document,
                "metadata": {
                    "source": "LaborLawNavigator",
                    "source_file": file_path.name,
                    "category": category,
                    "chunk_index": chunk_index,
                    "jurisdiction": "全国通用/非重庆专属",
                    "source_kind": "curated_local_dataset",
                },
            })
    return records, skipped


def should_include_file(file_path: Path) -> Tuple[bool, str]:
    name = file_path.name
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        return False, f"unsupported suffix {suffix}"
    if any(keyword in name for keyword in EXCLUDE_NAME_KEYWORDS):
        return False, "excluded by filename"
    if not any(keyword in name for keyword in USEFUL_KEYWORDS):
        return False, "not relevant enough by filename"
    return True, "selected"


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".docx" and docx is not None:
        document = docx.Document(str(file_path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    if suffix == ".pdf" and PyPDF2 is not None:
        parts = []
        with file_path.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages[:120]:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    if suffix == ".xlsx" and pd is not None:
        sheets = pd.read_excel(file_path, sheet_name=None)
        return "\n".join(frame.to_string(index=False) for frame in sheets.values())
    if suffix == ".csv" and pd is not None:
        return pd.read_csv(file_path).to_string(index=False)
    if suffix in {".txt", ".json", ".doc"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    return ""


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text.strip()


def chunk_text(text: str, size: int = 900, overlap: int = 120, max_chunks: int = 36) -> Iterable[str]:
    start = 0
    count = 0
    while start < len(text) and count < max_chunks:
        end = min(len(text), start + size)
        chunk = text[start:end].strip()
        if len(chunk) >= 120:
            yield chunk
            count += 1
        if end == len(text):
            break
        start = max(0, end - overlap)


def categorize(name: str, sample: str) -> str:
    haystack = f"{name} {sample}"
    for category, keywords in CATEGORY_RULES:
        if any(keyword in haystack for keyword in keywords):
            return category
    return "other"


def format_document(source_file: str, category: str, chunk_index: int, chunk: str) -> str:
    return (
        f"资料来源: LaborLawNavigator\n"
        f"文件名: {source_file}\n"
        f"分类: {category}\n"
        f"片段: {chunk_index + 1}\n"
        f"内容摘要片段: {chunk}"
    )


def stable_id(source_file: str, chunk_index: int) -> str:
    digest = hashlib.sha1(source_file.encode("utf-8")).hexdigest()[:16]
    return f"lln_{digest}_{chunk_index:03d}"


if __name__ == "__main__":
    main()
